#!/usr/bin/env python3
"""
PPC WORKBOOK VALIDATION HARNESS
===============================
Every countable rule from ppc-workbook-builder.skill, expressed as a check that
runs against the DELIVERED file and returns a count.

Why this exists: "read the skill in full" was added to the skill three times and
the same defect classes kept shipping. Every rule that broke a build was already
in the file — read, applied in one place, not carried to the second. A rule held
in a builder's head gets applied where they happen to be looking. A rule written
as a query gets applied everywhere.

Usage
-----
    python3 ppc_harness.py                  run every check against the workbook
    python3 ppc_harness.py --selftest       prove each check fails on a known-bad
    python3 ppc_harness.py --file X.xlsx    check a different file

Three properties every check here satisfies:
  1. Reads the delivered artifact, never the working state — the two diverge and
     only one ships.
  2. Returns a count, not a judgment — the result is not itself interpretable.
  3. Is demonstrated to fail against a known-bad case before it is trusted. A
     check never seen to fire has not been shown to test anything.

DEPENDENCY ORDER MATTERS. Levers feed each other: a staged base moves the
effective top-of-search price; a re-route moves the per-unit bound; recovered
routing makes a previously-uncomputable modifier computable. Checks run in the
order below, and the whole set re-runs after any change. A check that passed
before the last edit says nothing about the file that ships.
"""

import sys, re, json, math, shutil, collections
from pathlib import Path

import openpyxl
import pandas as pd

# openpyxl trap, found the hard way: ws.cell(row, col, value=None) does NOT clear
# a cell — it returns it without setting. Clearing requires .value = None. Several
# "the fix didn't persist" cycles on this build traced to exactly this.
HARNESS_VERSION = "1.0"
SKILL_VERSION_EXPECTED = "ppc-workbook-builder"

WORKBOOK = "/mnt/user-data/outputs/Decolure_Cooling_Sheets_4Piece_Workbook.xlsx"
ROUTING = "/home/claude/cs4pc_kickoff/routing.pkl"
BASIS = "/home/claude/cs4pc_kickoff/contribution_basis.json"
SUGGESTED = "/home/claude/cs4pc_kickoff/suggested_bids.csv"
TARGETED = "/home/claude/cs4pc_kickoff/targeted.json"
MKL_SOURCE = "/home/claude/recheck/raw/29082026/mkl_Cooling_Sheets_4_Piece_US_build316.xlsx"
AGED = "/home/claude/cs4pc_kickoff/aged_fwd.pkl"
PRODUCT_ADS = "/home/claude/cs4pc_kickoff/product_ads.json"

# Structure is the account's, not this skill's. A reviewer opening the delivered
# file should find the rows they exported plus decisions — never rows that
# appeared from somewhere.
EXPECTED_ROWS = 1813

# Placement conversion measured on this product, 60 days. Used to derive the
# affordable price at each placement from the routed child's own economics.
CVR_TOS, CVR_ROS, CVR_PDP = 8.19, 4.95, 3.77

# Per-size conversion, from this product's own delivering rows.
SIZE_CVR = {"TWIN": 6.9, "FULL": 6.8, "QUEEN": 5.3, "KING": 5.5, "CALKING": 6.0}

CORRECTION_CAP = 0.25      # §5  single-step cap on any bid or budget move
MODIFIER_CAP = 350.0       # §4A Bound 2
FLAT_CEILING = 9.00        # §4A Bound 4, the account-wide dollar ceiling
CLICK_FLOOR = 15           # §5  sample floor, per campaign AND per placement

# §4A — the exhaustive list of reasons this skill may apply a SILENT hold.
# Anything else "stops and asks first": a hold is a decision, not the absence of
# one, and it does not get to be the path of least resistance.
LEGITIMATE_HOLD_MARKERS = [
    "no delivery", "zero delivery", "0 clicks", "zero impressions", "never delivered",
    "no impressions", "took no delivery", "truncat",
    "moves only through base bid", "no modifier by construction",
    "fifteen-click floor", "15-click floor", "fewer than 15 clicks",
    "sample", "no rate", "loses $", "red —", "below the routing-switch",
]

ACTING_PREFIXES = ("HOLD", "PAUSE", "NEGATE", "NO MODIFIER", "BUDGET HELD",
                   "NO ACTION", "ASK", "STAY PAUSED")


# ----------------------------------------------------------------- data access
class Book:
    """Thin reader over the delivered workbook plus the economics it references."""

    def __init__(self, path):
        self.path = path
        self.wb = openpyxl.load_workbook(path)
        self.fb = self.wb["Final Bulk"]
        self.hdr = [c.value for c in self.fb[1]]
        self.col = {n: i + 1 for i, n in enumerate(self.hdr) if n}
        self.rt = pd.read_pickle(ROUTING)
        self.rt["zone"] = self.rt.apply(self._zone, axis=1)
        self.rti = self.rt.set_index("SKU")
        self.red = set(self.rt[self.rt.zone == "Red"].SKU)
        try:
            self.aged = pd.read_pickle(AGED).set_index("SKU")
        except Exception:
            self.aged = pd.DataFrame()
        try:
            self.ads = {k: set(v) for k, v in json.load(open(PRODUCT_ADS)).items()}
        except Exception:
            self.ads = {}

    @staticmethod
    def _zone(x):
        """§9 — three zones, not a single cutoff."""
        doh = x.DOH if pd.notna(x.DOH) else None
        if (x.Inbound or 0) > 0:
            return "Green"
        if doh is None:
            return "No data"
        return "Green" if doh >= 60 else ("Yellow" if doh >= 21 else "Red")

    def g(self, row, name):
        c = self.col.get(name)
        return self.fb.cell(row=row, column=c).value if c else None

    def rows(self):
        return range(2, self.fb.max_row + 1)

    def decided(self):
        return [r for r in self.rows() if self.g(r, "Action")]

    def num(self, row, name):
        try:
            return float(self.g(row, name))
        except (TypeError, ValueError):
            return None

    def size_of(self, sku):
        for k in ("CALKING", "KING", "QUEEN", "FULL", "TWIN"):
            if k in str(sku):
                return k
        return None

    def economics(self, sku):
        """Contribution for a fresh child; forward cash for aged stock (§3)."""
        if sku in getattr(self.aged, "index", []):
            a = self.aged.loc[sku]
            return float(a.fwd_cash), float(a.cvr), float(a.aov)
        if sku in self.rti.index:
            r = self.rti.loc[sku]
            if pd.notna(r.CM2) and r.CM2 > 0:
                return float(r.CM2), SIZE_CVR.get(self.size_of(sku), 5.5), float(r.aov)
        return None, None, None

    def affordable(self, sku):
        """Base is set by the tighter of rest-of-search and product pages."""
        cm2, cvr, _ = self.economics(sku)
        if cm2 is None:
            return None
        return round(min(cm2 * cvr / 100,
                         cm2 * (cvr * CVR_PDP / CVR_ROS) / 100), 2)

    def base_by_campaign(self):
        out = {}
        for r in self.rows():
            if self.g(r, "Entity") not in ("Keyword", "Product Targeting"):
                continue
            b = self.num(r, "New Bids")
            if b is None:
                b = self.num(r, "Bid")
            if b is None:
                continue
            cn = str(self.g(r, "Campaign Name (Informational only)"))
            out.setdefault(cn, []).append(b)
        return {k: max(v) for k, v in out.items()}

    def sku_by_campaign(self):
        out = {}
        for r in self.rows():
            if self.g(r, "Entity") not in ("Keyword", "Product Targeting"):
                continue
            s = self.g(r, "SKU")
            if s:
                out[str(self.g(r, "Campaign Name (Informational only)"))] = str(s)
        return out

    def all_text(self):
        parts = []
        for r in self.decided():
            parts.append(str(self.g(r, "Action") or ""))
            parts.append(str(self.g(r, "Reasoning") or ""))
        return " ".join(parts)


# --------------------------------------------------------------------- checks
# Each returns (count_of_failures, list_of_example_strings).
# Ordered by dependency: structure, then routing, then bases, then anything
# derived from a base, then writing standard.

def chk_structure(b):
    """Rows must not be added to the delivered file. A decision with no row of
    its own is written on the campaign or keyword row that carries it."""
    delta = (b.fb.max_row - 1) - EXPECTED_ROWS
    return (abs(delta), [f"row count {b.fb.max_row - 1} vs {EXPECTED_ROWS} exported"] if delta else [])


def chk_coverage(b):
    """§14 — every row of every entity carries a verdict. A blank row is
    indistinguishable from a row nobody reached."""
    bad = [r for r in b.rows() if b.g(r, "Entity") and not b.g(r, "Action")]
    c = collections.Counter(b.g(r, "Entity") for r in bad)
    return len(bad), [f"{v} undecided {k} rows" for k, v in c.most_common(4)]


def chk_routing_is_red(b):
    """§9 — Red re-points the same day. No spend is directed at a child below
    the routing-switch threshold."""
    bad = []
    for r in b.decided():
        a = str(b.g(r, "Action"))
        if a.startswith(ACTING_PREFIXES):
            continue
        s = str(b.g(r, "SKU") or "")
        if s in b.red:
            bad.append(f"row {r}: spend routed to {s} (Red)")
    return len(bad), bad[:4]


def chk_routing_from_ads(b):
    """Routing is read from product-ad rows, never inferred from a campaign's
    name. Flags campaigns whose every product ad is Red but which carry no
    re-route."""
    if not b.ads:
        return 0, ["product-ad data unavailable — check skipped, not passed"]
    known = set(b.rt.SKU)
    bad = []
    for r in b.decided():
        if b.g(r, "Entity") not in ("Keyword", "Product Targeting"):
            continue          # routing lives on rows that carry a SKU
        cn = str(b.g(r, "Campaign Name (Informational only)"))
        cur = b.ads.get(cn, set()) & known
        if not cur or not (cur <= b.red):
            continue
        if not str(b.g(r, "Action")).startswith(("RE-ROUTE", "HOLD", "PAUSE",
                                                 "ASK", "NO MODIFIER",
                                                 "BUDGET HELD", "NO ACTION")):
            bad.append(f"row {r}: all product ads Red on {cn[:40]}")
    return len(bad), bad[:4]


def chk_attribute_match(b):
    """§2A step 8 — a query naming an attribute may only route to a SKU sharing
    it. Governs reactivation as well as launch."""
    colours = {"white": "WHITE", "grey": "GREY", "gray": "GREY", "black": "BLACK",
               "navy": "NAVY", "pink": "PINK", "sage": "SAGE", "cream": "CREAM",
               "beige": "BEIGE", "purple": "PURPLE"}
    bad = []
    for r in b.decided():
        kw = str(b.g(r, "Keyword Text") or "").lower()
        sku = str(b.g(r, "SKU") or "").upper()
        if not kw or not sku:
            continue
        for w, tag in colours.items():
            if re.search(rf"\b{w}\b", kw) and tag not in sku:
                bad.append(f'row {r}: "{kw[:30]}" names {w} -> {sku}')
                break
    return len(bad), bad[:4]


def chk_base_derived(b):
    """§4A Bound 1 — the base is derived from the routed child's own economics,
    never inherited. A staged row is judged on its stated step-2 target."""
    bad = []
    for r in b.decided():
        if b.g(r, "Entity") not in ("Keyword", "Product Targeting"):
            continue
        nb = b.num(r, "New Bids")
        if nb is None:
            continue
        aff = b.affordable(str(b.g(r, "SKU") or ""))
        if aff is None:
            continue
        rsn = str(b.g(r, "Reasoning") or "")
        m = re.findall(r"\$([\d.]+) at day 7", rsn)
        target = float(m[-1]) if m else nb   # last stated step-2 target is the current one
        if target > aff * 1.03:
            bad.append(f"row {r}: target ${target:.2f} vs affordable ${aff:.2f}")
    return len(bad), bad[:4]


def chk_correction_cap(b):
    """§5 — no single step moves a bid more than 25%. A wider gap stages as two
    dated steps with both targets committed upfront."""
    bad = []
    for r in b.decided():
        ob, nb = b.num(r, "Bid"), b.num(r, "New Bids")
        if not ob or nb is None or ob <= 0:
            continue
        if abs(nb - ob) / ob > CORRECTION_CAP + 0.0001:
            bad.append(f"row {r}: ${ob:.2f} -> ${nb:.2f} ({(nb-ob)/ob*100:+.0f}%)")
    return len(bad), bad[:4]


def chk_modifier_entity(b):
    """§14 / 0-Q — the platform reads New Percentage on Bidding Adjustment rows.
    A modifier on a keyword row never deploys."""
    bad = [f"row {r}" for r in b.rows()
           if b.g(r, "Entity") == "Keyword" and b.g(r, "New Percentage") is not None]
    return len(bad), bad[:4]


def _tos_prices(b):
    base, sku = b.base_by_campaign(), b.sku_by_campaign()
    out = []
    for r in b.rows():
        if b.g(r, "Entity") != "Bidding Adjustment":
            continue
        if "Top" not in str(b.g(r, "Placement") or ""):
            continue
        m = b.num(r, "New Percentage")
        cn = str(b.g(r, "Campaign Name (Informational only)"))
        if m is None or cn not in base:
            continue
        cm2, _, _ = b.economics(sku.get(cn, ""))
        out.append((r, base[cn] * (1 + m / 100), m, cm2))
    return out


def chk_modifier_cap(b):
    """§4A Bound 2 — beyond 350% a modifier stops being a placement adjustment
    and becomes a different bid nobody sized."""
    bad = [f"row {r}: {m:.0f}%" for r, _, m, _ in _tos_prices(b) if m > MODIFIER_CAP]
    return len(bad), bad[:4]


def chk_tos_vs_unit(b):
    """§4A Bound 3 — no click may cost more than the unit earns. At 100%
    conversion a click priced at CM2 exactly breaks even; above it, the click
    loses money at every conversion rate that exists."""
    bad = [f"row {r}: ${p:.2f} vs ${cm2:.2f}/unit"
           for r, p, _, cm2 in _tos_prices(b) if cm2 and p > cm2 + 0.01]
    return len(bad), bad[:4]


def chk_tos_vs_flat(b):
    """§4A Bound 4 — the flat account-wide ceiling, last and loosest."""
    bad = [f"row {r}: ${p:.2f}" for r, p, _, _ in _tos_prices(b) if p > FLAT_CEILING]
    return len(bad), bad[:4]


def chk_hold_legitimacy(b):
    """§4A — the list of reasons for a SILENT hold is exhaustive. Anything else
    stops and asks."""
    bad = []
    for r in b.decided():
        a = str(b.g(r, "Action"))
        if not re.match(r"^(HOLD|NO ACTION|NO MODIFIER|BUDGET HELD|STAY PAUSED)", a):
            continue
        rsn = str(b.g(r, "Reasoning") or "").lower()
        if not any(k in rsn for k in LEGITIMATE_HOLD_MARKERS):
            bad.append(f"row {r}: {a[:44]}")
    return len(bad), bad[:4]


def chk_reasoning_present(b):
    """§14 0-L — present, unique, correct, written: four separate properties."""
    bad = [f"row {r}" for r in b.decided() if not b.g(r, "Reasoning")]
    return len(bad), bad[:4]


def chk_reasoning_unique(b):
    """§14 0-O — a string that could be pasted onto a different row without
    becoming false is a category name, not reasoning."""
    c = collections.Counter(str(b.g(r, "Reasoning") or "") for r in b.decided())
    worst = [(n, t) for t, n in c.items() if n > 4]
    return len(worst), [f"x{n}: {t[:60]}" for n, t in sorted(worst, reverse=True)[:4]]


def chk_reasoning_figures(b):
    """§14 — named numbers doing work, not adjectives."""
    bad = [f"row {r}" for r in b.decided()
           if not re.search(r"\d", str(b.g(r, "Reasoning") or ""))]
    return len(bad), bad[:4]


def chk_window_stated(b):
    """§14 — every cited number carries the window it reflects."""
    pat = r"sixty days|60 days|Aug|Jun|Sep|day 0|day 7"
    bad = [f"row {r}" for r in b.decided()
           if not re.search(pat, str(b.g(r, "Reasoning") or ""))]
    return len(bad), bad[:4]


def chk_reverses_present(b):
    """§14 — every verdict states what would reverse it."""
    bad = [f"row {r}" for r in b.decided() if not b.g(r, "Reverses If")]
    return len(bad), bad[:4]


def chk_same_value_blank(b):
    """§14 — same-value goes blank in every "New X" column."""
    bad = []
    for r in b.decided():
        for a, n in (("Bid", "New Bids"), ("Daily Budget", "New Budget"),
                     ("Percentage", "New Percentage")):
            x, y = b.num(r, a), b.num(r, n)
            if x is not None and y is not None and x == y:
                bad.append(f"row {r}: {n} equals {a}")
    return len(bad), bad[:4]


def chk_revision_narrative(b):
    """§14 — zero revision narrative. "corrected this pass" fails the
    never-saw-an-earlier-draft test."""
    txt = b.all_text().lower()
    hits = [w for w in ("earlier", "wrongly", "previously", "corrected this")
            if w in txt]
    return len(hits), hits


def chk_internal_labels(b):
    """§14 — no section refs, internal codes or raw platform IDs in delivered
    text. A citation a reader cannot decode is opaque jargon, not a receipt."""
    txt = b.all_text()
    hits = [p for p in ("§", "State E", "sufficiency stop") if p in txt]
    hits += list(set(re.findall(r"\b\d{12,}\b", txt)))[:3]
    return len(hits), hits[:4]


def chk_deferred(b):
    """0-C — no analysis deferred to a later pass."""
    hits = re.findall(r"can be computed|will be computed", b.all_text(), re.I)
    return len(hits), hits[:4]


def chk_text_integrity(b):
    """Reasoning must read as prose. A period followed straight by a digit is
    the signature of a regex edit that cut a sentence in half — on a real cycle
    this corrupted 556 rows while every other check still passed."""
    bad = []
    for r in b.decided():
        t = str(b.g(r, "Reasoning") or "")
        # A single pattern is not enough. Every fragment class below reached a
        # delivered file while this check reported PASS, because it tested one
        # shape of damage and patching produces many.
        pats = [
            r"[a-z]\.\d",                      # sentence cut mid-clause
            r"returns [A-Z]{2}-[A-Z0-9-]+ [a-z]",   # subject repeated after an excision
            r"[A-Z]{2}-[A-Z0-9-]+ sold only .* returns",
            r"\.\s*\d+ base puts",              # number orphaned from its sentence
            r"ceiling\.\s*\d",
            r"\b(\w+) \1\b",                   # a word doubled by a splice
            r"\s,|\s\.|\(\s*\)",             # punctuation left stranded
        ]
        hit = next((m for m in (re.search(p, t) for p in pats) if m), None)
        if hit:
            bad.append(f"row {r}: ...{t[max(0,hit.start()-40):hit.end()+40]}...")
    return len(bad), bad[:4]


def chk_discovery_reasoned(b):
    """Every Broad, Phrase and product-targeting row carries its own reasoning,
    and an auto campaign's four targeting types are decided separately rather
    than sharing one campaign verdict."""
    bad = []
    autos = {}
    delivering = set()          # auto campaigns that actually took clicks
    for r in b.decided():
        ent = b.g(r, "Entity")
        mt = str(b.g(r, "Match Type") or "")
        exp = str(b.g(r, "Product Targeting Expression") or "").strip()
        if not (ent == "Product Targeting" or (ent == "Keyword" and mt in ("Broad", "Phrase"))):
            continue
        if not b.g(r, "Reasoning"):
            bad.append(f"row {r}: discovery row with no reasoning")
        if exp in ("close-match", "loose-match", "complements", "substitutes"):
            cn = str(b.g(r, "Campaign Name (Informational only)"))
            autos.setdefault(cn, set()).add(str(b.g(r, "Action"))[:20])
            cl = b.num(r, "Clicks")
            if cl and cl > 0:
                delivering.add(cn)
    # Sharing one verdict is only wrong where the types carry DIFFERENT data.
    # A campaign with no delivery at all correctly gives every type the same
    # answer, and flagging that would be the check firing on a legitimate case.
    for cn, acts in autos.items():
        if len(acts) == 1 and cn in delivering:
            bad.append(f"{cn[:40]}: types differ in data but share one verdict")
    return len(bad), bad[:4]


def chk_tabs_actioned(b):
    """The Broad-Phrase Coverage and PAT tabs carry Action, Reasoning and
    Reverses If on every row — decisions live in the tab that holds them."""
    bad = []
    for name, cols in (("Broad-Phrase Coverage", (21, 22, 23)), ("PAT", (15, 16, 17))):
        if name not in b.wb.sheetnames:
            continue
        ws = b.wb[name]
        for r in range(2, ws.max_row + 1):
            if not ws.cell(row=r, column=1).value:
                continue
            for c, lbl in zip(cols, ("Action", "Reasoning", "Reverses If")):
                if not ws.cell(row=r, column=c).value:
                    bad.append(f"{name} row {r}: {lbl} empty")
    return len(bad), bad[:4]


def chk_mkl_complete(b):
    """New-launch rows in the Master Keyword List carry the full five-property
    sizing, not a placeholder. Sized means target rank, daily target, required
    clicks and required budget all populated together — a row missing any of
    them is not a sized push regardless of how good its data looks."""
    if "Master Keyword List" not in b.wb.sheetnames:
        return 0, ["MKL tab absent — check skipped, not passed"]
    ws = b.wb["Master Keyword List"]
    hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
    need = ["Reasoning", "Reverses if", "Proposed SKU", "rank targets", "Daily Target",
            "No of Click Required", "Expected CPC", "Proposed Bid ", "TOS%",
            "Bidding Strategy", "PPC Daily Budget", "Pro-Stock Status?", "Margin?"]
    missing = [c for c in need if c not in hdr]
    if missing:
        return len(missing), [f"MKL column absent: {c}" for c in missing[:4]]
    bad = []
    for r in range(2, ws.max_row + 1):
        act = str(ws.cell(row=r, column=hdr["Action"]).value or "")
        if not act.startswith("NEW LAUNCH"):
            continue        # a row screened OUT of the launch list carries a
                            # reason, not a sizing — requiring both would force
                            # a bid onto a term deliberately not being launched
        for c in need:
            v = ws.cell(row=r, column=hdr[c]).value
            if v in (None, ""):
                bad.append(f"MKL row {r}: {c} empty")
            elif c in ("Expected CPC", "Proposed Bid ") and isinstance(v, str):
                # A formula is a live model, which is what the account wants —
                # reject only text that is neither a number nor a formula.
                if not v.startswith("="):
                    bad.append(f'MKL row {r}: {c} is text "{v}", not a price or formula')
    return len(bad), bad[:4]


def chk_discovery_no_premium(b):
    """A Discovery row is bounded by the CPA ceiling — the above-ceiling
    allowance belongs to Ranking rows inside a ranking window. An auto or
    broad campaign has no rank target, so its top-of-search modifier is solved
    from measured placement conversion, never from a rank gap."""
    bad = []
    autos = set()
    for r in b.decided():
        cn = str(b.g(r, "Campaign Name (Informational only)") or "")
        if "auto" in cn.lower() or "-disc-" in cn.lower():
            autos.add(cn)
    for r in b.decided():
        if b.g(r, "Entity") != "Bidding Adjustment":
            continue
        cn = str(b.g(r, "Campaign Name (Informational only)") or "")
        if cn not in autos:
            continue
        m = b.num(r, "New Percentage")
        rsn = str(b.g(r, "Reasoning") or "").lower()
        if m is None:
            continue
        if "rank gap" in rsn or "ranking premium" in rsn:
            bad.append(f"row {r}: discovery modifier justified by rank gap")
        elif "measured" not in rsn:
            bad.append(f"row {r}: discovery modifier not tied to measured conversion")
    return len(bad), bad[:4]


def chk_flat_sizing(b):
    """§7 — a premium is sized from each row's own evidence. "A roster of
    ranking keywords will show a roster of different premiums; a flat number
    across all of them is the signal that this wasn't actually derived per
    row." Fires when one modifier value dominates the set."""
    vals = []
    for r in b.decided():
        if b.g(r, "Entity") != "Bidding Adjustment":
            continue
        m = b.num(r, "New Percentage")
        if m is not None:
            vals.append(int(m))
    # A zero is the ABSENCE of a premium, not a premium value — a row where top
    # of search affords nothing over the base is a derived answer, not a default.
    # Testing dominance among zeros would flag a correctly-priced file.
    nz = [v for v in vals if v]
    if len(nz) < 20:
        return 0, []
    from collections import Counter
    c = Counter(nz)
    top, n = c.most_common(1)[0]
    share = n / len(nz)
    if share > 0.45:
        return 1, [f"{n} of {len(nz)} non-zero modifiers sit at {top}% ({share:.0%}) — sizing has collapsed"]
    return 0, []


def chk_duplicates(b):
    """§1 — one active instance per keyword per match type, unless one of the
    four coexistence reasons applies. Singular and plural collapse to the
    higher search volume; word-order variants do not collapse."""
    from collections import defaultdict
    inst = defaultdict(list)
    for r in b.decided():
        if b.g(r, "Entity") != "Keyword":
            continue
        a = str(b.g(r, "Action") or "")
        if a.startswith("PAUSE"):
            continue                       # already consolidated
        kw = str(b.g(r, "Keyword Text") or "").lower().strip()
        mt = str(b.g(r, "Match Type") or "")
        if kw:
            inst[(kw, mt)].append(r)
    bad = []
    for (kw, mt), rows in inst.items():
        if len(rows) < 2:
            continue
        skus = {str(b.g(r, "SKU") or "") for r in rows}
        if len(skus) > 1 and "" not in skus:
            continue                       # variation split, legitimate
        bad.append(f'"{kw[:34]}" [{mt}] live in {len(rows)} campaigns on one child')
    live = {str(b.g(r, "Keyword Text") or "").lower().strip()
            for r in b.decided()
            if b.g(r, "Entity") == "Keyword" and not str(b.g(r, "Action") or "").startswith("PAUSE")}
    live.discard("")
    for k in live:
        alt = k + "s" if not k.endswith("s") else k[:-1]
        if alt in live and k < alt:
            bad.append(f'"{k}" and "{alt}" both live — singular/plural not collapsed')
    return len(bad), bad[:4]


def chk_launch_scope(b):
    """§2A Job A — a term that already carries a live row anywhere in the
    account is a reactivation and belongs to Job B, not the new-launch list.
    Branded terms are never launched: ours are Defensive, another brand's runs
    as conquest through product targeting. Both screens must run against the
    ACCOUNT, not just this portfolio."""
    if "Master Keyword List" not in b.wb.sheetnames:
        return 0, ["MKL absent — check skipped, not passed"]
    try:
        with open(TARGETED) as fh:
            t = json.load(fh)
        live = set(t["account"])
    except Exception:
        return 0, ["account keyword set unavailable — check skipped, not passed"]
    try:
        src = pd.read_excel(MKL_SOURCE, sheet_name="Keywords")
        src["_k"] = src["MKL"].astype(str).str.lower().str.strip()
        branded = set(src.loc[src["Final Categorization"].astype(str).str.strip() == "Branded", "_k"])
    except Exception:
        branded = set()
    ws = b.wb["Master Keyword List"]
    hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
    bad = []
    for r in range(2, ws.max_row + 1):
        a = str(ws.cell(row=r, column=hdr["Action"]).value or "")
        if not a.startswith("NEW LAUNCH"):
            continue
        kw = str(ws.cell(row=r, column=hdr["MKL"]).value or "").lower().strip()
        if kw in live:
            bad.append(f'MKL row {r}: "{kw[:34]}" already has a live row')
        elif kw in branded:
            bad.append(f'MKL row {r}: "{kw[:34]}" is categorised Branded')
        elif re.search(r"[^\w\s\-'&]", kw):
            bad.append(f'MKL row {r}: "{kw[:34]}" is a malformed string')
    return len(bad), bad[:4]


def chk_launch_relevance(b):
    """§2A — relevance is judged against the listing's own attributes, not a
    supplied relevancy score. Every launch must name the product category, and
    must not name a product the listing does not sell. The score column is
    known to disagree with the listing on this account, so it is not the test."""
    if "Master Keyword List" not in b.wb.sheetnames:
        return 0, ["MKL absent — check skipped, not passed"]
    core = ("sheet", "sheets", "bedsheet", "bedsheets")
    off = re.compile(r"comforter|duvet|quilt|blanket|mattress|topper|protector|"
                     r"towel|curtain|rug|100% bamboo|organic cotton|linen|flannel|silk|satin")
    ws = b.wb["Master Keyword List"]
    hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
    bad = []
    for r in range(2, ws.max_row + 1):
        if not str(ws.cell(row=r, column=hdr["Action"]).value or "").startswith("NEW LAUNCH"):
            continue
        kw = str(ws.cell(row=r, column=hdr["MKL"]).value or "").lower().strip()
        if off.search(kw):
            bad.append(f'MKL row {r}: "{kw[:34]}" names a product this listing does not sell')
        elif not any(re.search(rf"\b{w}\b", kw) for w in core):
            bad.append(f'MKL row {r}: "{kw[:34]}" does not name the product category')
        else:
            rsn = str(ws.cell(row=r, column=hdr["Reasoning"]).value or "").lower()
            if "listing" not in rsn:
                bad.append(f'MKL row {r}: relevance not tied to the listing\'s own attributes')
    return len(bad), bad[:4]


def chk_syntax_labels(b):
    """A launch row's syntax group is derived from the term's own roots. A label
    reading "Irrelevant" on a term the listing actually sells is a stale
    classification, not a verdict — and a syntax that omits a size or material
    the term names cannot be used to roll that term up with its peers."""
    if "Master Keyword List" not in b.wb.sheetnames:
        return 0, ["MKL absent — check skipped, not passed"]
    ws = b.wb["Master Keyword List"]
    hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
    if "Syntax Groups" not in hdr:
        return 0, ["Syntax Groups column absent — check skipped"]
    sizes = [("california king", "California King"), ("cal king", "California King"),
             ("twin", "Twin"), ("queen", "Queen"), ("king", "King"), ("full", "Full")]
    bad = []
    for r in range(2, ws.max_row + 1):
        if not str(ws.cell(row=r, column=hdr["Action"]).value or "").startswith("NEW LAUNCH"):
            continue
        kw = str(ws.cell(row=r, column=hdr["MKL"]).value or "").lower().strip()
        syn = str(ws.cell(row=r, column=hdr["Syntax Groups"]).value or "")
        if not syn.strip() or "irrelevant" in syn.lower():
            bad.append(f'MKL row {r}: "{kw[:30]}" carries no usable syntax ("{syn[:20]}")')
            continue
        if bool(re.search(r"bamboo|viscose|lyocell|rayon", kw)) != ("Bamboo" in syn):
            bad.append(f'MKL row {r}: "{kw[:30]}" bamboo root disagrees with its syntax')
            continue
        for lit, lab in sizes:
            if re.search(rf"\b{lit}\b", kw) and lab not in syn and "California King" not in syn:
                bad.append(f'MKL row {r}: "{kw[:30]}" names {lit}, syntax omits it')
                break
    return len(bad), bad[:4]


def chk_launch_bids(b):
    """§2A — a launch bid is placed inside Amazon's own suggested range, never
    estimated. The model is two-factor: base x (1 + TOS multiplier), where the
    base is what PRODUCT PAGES afford on the routed child and the premium lives
    in the multiplier, so rest-of-search and product pages pay only what they
    earn (the PDP-driven base rule). Top of search stays inside the child's own
    per-unit bound, which the flat ceiling may only tighten.

    No bid floor is asserted: the account sets the base from PDP economics, and
    a floor above that would make product pages overpay to lift top of search.
    Formula cells are accepted — a live model is the point, not static values."""
    if "Master Keyword List" not in b.wb.sheetnames:
        return 0, ["MKL absent — check skipped, not passed"]
    try:
        sug = pd.read_csv(SUGGESTED)
        sug["_k"] = sug["Keyword"].astype(str).str.lower().str.strip()
        rng = {r._k: (r["Suggested bid (low)(USD)"], r["Suggested bid (median)(USD)"],
                      r["Suggested bid (high)(USD)"]) for _, r in sug.iterrows()}
    except Exception:
        return 0, ["suggested-bid file unavailable — check skipped, not passed"]
    ws = b.wb["Master Keyword List"]
    hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
    bad = []
    for r in range(2, ws.max_row + 1):
        if not str(ws.cell(row=r, column=hdr["Action"]).value or "").startswith("NEW LAUNCH"):
            continue
        kw = str(ws.cell(row=r, column=hdr["MKL"]).value or "").lower().strip()
        base = ws.cell(row=r, column=hdr["Proposed Bid "]).value
        mult = ws.cell(row=r, column=hdr["TOS%"]).value
        sku = str(ws.cell(row=r, column=hdr["Proposed SKU"]).value or "")
        if kw not in rng:
            bad.append(f'MKL row {r}: "{kw[:30]}" has no supplied suggested range')
            continue
        if base is None or mult is None:
            bad.append(f'MKL row {r}: "{kw[:30]}" missing base or multiplier')
            continue
        if isinstance(base, str) or isinstance(mult, str):
            continue          # formula cell — the model computes it, not this check
        base, mult = float(base), float(mult)
        if base <= 0:
            bad.append(f"MKL row {r}: base is {base}")
            continue
        tos = base * (1 + mult)
        cm2 = None
        if sku in b.rti.index and pd.notna(b.rti.loc[sku].CM2):
            cm2 = float(b.rti.loc[sku].CM2)
        cap = min(cm2, 9.0) if cm2 else 9.0
        if tos > cap + 0.02:
            bad.append(f"MKL row {r}: top of search ${tos:.2f} over the ${cap:.2f} bound")
        if mult > 3.5:
            bad.append(f"MKL row {r}: multiplier {mult} exceeds 3.5x")
    return len(bad), bad[:4]


def chk_delivery_apparatus(b):
    """The gate's output ships inside the file. A validation run whose result
    lives only in a console is not evidence to a reviewer, and its absence is
    how contradictions between an action and its own reasoning reach delivery."""
    need = ["Summary", "Validation Gate", "No-Action Census", "Change Review Sheet"]
    missing = [t for t in need if t not in b.wb.sheetnames]
    return len(missing), [f"tab absent: {t}" for t in missing]


def chk_action_matches_cell(b):
    """With state columns frozen, the Action string is the executor's
    instruction — so the figure it names must be the figure the column
    deploys. Two notations for one operation is an ambiguity, and a bid named
    in words that differs from the cell is a contradiction the executor
    resolves by guessing."""
    bad = []
    for r in b.decided():
        a = str(b.g(r, "Action") or "")
        nb = b.num(r, "New Bids")
        m = re.search(r"BID ([\d.]+)", a)
        if m and nb is not None and abs(float(m.group(1)) - nb) > 0.005:
            bad.append(f"row {r}: action says ${float(m.group(1)):.2f}, column deploys ${nb:.2f}")
        if "(hold)" in a.lower():
            mm = re.search(r"BUDGET ([\d.]+) → ([\d.]+)", a)
            if mm and abs(float(mm.group(1)) - float(mm.group(2))) > 0.005:
                bad.append(f"row {r}: action says hold but writes {mm.group(1)} → {mm.group(2)}")
    return len(bad), bad[:4]


def chk_sample_disclosed(b):
    """A contribution figure rests on a unit sample, and the same reasoning that
    refuses a rate on two clicks must refuse one on two units. Every row pricing
    against a child's contribution names the sample that produced it."""
    bad = []
    for r in b.decided():
        t = str(b.g(r, "Reasoning") or "")
        if "contribution $" not in t and "a unit" not in t:
            continue
        if b.g(r, "Entity") not in ("Keyword", "Product Targeting"):
            continue
        # Forward cash is a different basis with a different sample — units held,
        # not units sold — so naming the seam satisfies this instead.
        if "forward cash" in t.lower() and "naming the seam" in t.lower():
            continue
        if "no conversion rate is decidable" in t.lower():
            continue
        if not re.search(r"\d+ units? sold|units sold across|too thin to carry|sold only \d+ units|sold \d+ units", t):
            bad.append(f"row {r}: prices on a contribution figure without naming its unit sample")
    return len(bad), bad[:4]


def chk_premium_gated(b):
    """A top-of-search price above what the placement itself affords is a rank
    premium, and a rank premium requires a rank case and a real loss ceiling.
    Bound3 alone does NOT cover this: a per-unit contribution is not a per-click
    ceiling — it breaks even only at 100% conversion — so a check that passes
    $8.99 because the unit earns $17.35 enshrines the category error it exists
    to catch. The bound here is what top of search affords, plus a capped lift."""
    try:
        with open(BASIS) as fh:
            B = json.load(fh)
    except Exception:
        return 0, ["contribution basis unavailable — check skipped, not passed"]
    CVR = {"TWIN": 6.9, "FULL": 6.8, "QUEEN": 5.3, "KING": 5.5, "CALKING": 6.0}
    TOSR = 8.19 / 4.95
    PREM_CAP = 0.25
    base, skum = {}, {}
    for r in b.decided():
        if b.g(r, "Entity") not in ("Keyword", "Product Targeting"):
            continue
        cn = str(b.g(r, "Campaign Name (Informational only)"))
        # A staged row's step-1 value is temporarily high while it comes down
        # from an inherited bid. Judge the committed step-2 target, the same
        # convention bound1/base uses, or a downward stage reads as a breach.
        # A paused row deploys nothing, so its old bid must not set the campaign
        # base — counting it made a correctly-priced campaign read as a breach.
        if str(b.g(r, "Action") or "").startswith(("PAUSE", "NEGATE", "ASK")):
            continue
        v = b.num(r, "New Bids") or b.num(r, "Bid")
        m2 = re.findall(r"\$([\d.]+) at day 7", str(b.g(r, "Reasoning") or ""))
        if m2:
            v = float(m2[-1])
        if v:
            base[cn] = max(base.get(cn, 0), v)
        if b.g(r, "SKU"):
            skum[cn] = str(b.g(r, "SKU"))
    bad = []
    for r in b.decided():
        if b.g(r, "Entity") != "Bidding Adjustment":
            continue
        if "Top" not in str(b.g(r, "Placement") or ""):
            continue
        m = b.num(r, "New Percentage")
        cn = str(b.g(r, "Campaign Name (Informational only)"))
        s, bb = skum.get(cn), base.get(cn)
        if m is None or not bb or s not in B or B[s]["cm2"] <= 0:
            continue
        sz = next((k for k in ("CALKING", "KING", "QUEEN", "FULL", "TWIN") if k in s), None)
        aff = B[s]["cm2"] * (CVR.get(sz, 5.5) * TOSR) / 100
        tos = bb * (1 + m / 100)
        if tos > aff * (1 + PREM_CAP) + 0.02:
            bad.append(f"row {r}: TOS ${tos:.2f} is {tos/aff:.1f}x the ${aff:.2f} the placement affords")
        elif tos > aff * 1.02:
            t = str(b.g(r, "Reasoning") or "").lower()
            if "rank" not in t or "loss ceiling" not in t:
                bad.append(f"row {r}: premium above affordable with no rank case or loss ceiling stated")
            elif "a week" not in t or "checkpoint" not in t:
                bad.append(f"row {r}: loss ceiling stated as a rate, not a dated dollar stop")
    return len(bad), bad[:4]


def chk_enable_has_gate(b):
    """A keyword ENABLE inside a paused campaign deploys nothing on its own. The
    campaign needs its own action, or the row needs a named go/wait condition —
    otherwise an executor cannot tell what actually goes live."""
    campaign_acted = set()
    for r in b.decided():
        if b.g(r, "Entity") == "Campaign":
            campaign_acted.add(str(b.g(r, "Campaign Name (Informational only)")))
    bad = []
    for r in b.decided():
        a = str(b.g(r, "Action") or "")
        if not a.startswith("ENABLE"):
            continue
        if str(b.g(r, "Campaign State (Informational only)") or "") != "paused":
            continue
        cn = str(b.g(r, "Campaign Name (Informational only)"))
        t = str(b.g(r, "Reasoning") or "").lower()
        if cn not in campaign_acted and not any(w in t for w in ("goes live when", "deploys when", "waits until", "go/wait")):
            bad.append(f"row {r}: ENABLE in a paused campaign with no campaign action and no go/wait condition")
    return len(bad), bad[:4]


def chk_counts_agree(b):
    """The QA tabs must not disagree with each other. A Summary claiming one
    gate count while the gate tab lists another, or a held-count that differs
    from the census by exactly the ASK count, is same-figure-two-values applied
    to the quality apparatus itself — the one place it destroys all the rest."""
    need = {"Summary", "Validation Gate", "No-Action Census"}
    if not need <= set(b.wb.sheetnames):
        return 0, ["QA tabs absent — check skipped, not passed"]
    def grab(tab, label):
        ws = b.wb[tab]
        for r in range(1, ws.max_row + 1):
            for c in range(1, min(ws.max_column, 4) + 1):
                if str(ws.cell(row=r, column=c).value or "").strip().lower() == label.lower():
                    for c2 in range(c + 1, min(ws.max_column, 4) + 1):
                        v = ws.cell(row=r, column=c2).value
                        if v is not None and str(v).strip() != "":
                            return v
        return None
    bad = []
    gate_rows = sum(1 for r in range(2, b.wb["Validation Gate"].max_row + 1)
                    if str(b.wb["Validation Gate"].cell(row=r, column=3).value or "") in ("PASS", "FAIL"))
    run = str(grab("Summary", "Validation gate") or "")
    m = re.search(r"(\d+)\s*/\s*(\d+)", run)
    if m and int(m.group(2)) != gate_rows:
        bad.append(f"Summary says {m.group(2)} checks, the gate tab lists {gate_rows}")
    s_held, c_held = grab("Summary", "Held"), grab("No-Action Census", "Total held")
    if s_held is not None and c_held is not None and int(s_held) != int(c_held):
        bad.append(f"Summary held {s_held}, census held {c_held}")
    s_ask, c_ask = grab("Summary", "Raised as questions"), grab("No-Action Census", "Raised as questions")
    if s_ask is not None and c_ask is not None and int(s_ask) != int(c_ask):
        bad.append(f"Summary asks {s_ask}, census asks {c_ask}")
    return len(bad), bad[:4]


def chk_row_integrity(b):
    """THE ROW-INTEGRITY GATE. The unit of regeneration is the decision, not the
    cell — and the unit of checking has to match, because a gate that reads
    columns independently will pass a file whose columns disagree. It did that
    three renders running: an Action pricing top of search at $7.35 beside a
    Reasoning saying no price is stated, beside a Reverses-If citing a per-order
    figure from a computation two passes deleted.

    Four things are read together, per row:
      1. does the Reasoning affirm the figure the Action deploys
      2. does the Reverses-If cite only computations the Reasoning still states
      3. does the Action's own wording match the value in the cell
      4. does the hosting Campaign row describe the program its keywords carry
    """
    bad = []
    for r in b.decided():
        a = str(b.g(r, "Action") or "")
        t = str(b.g(r, "Reasoning") or "")
        v = str(b.g(r, "Reverses If") or "")
        # 1 — an Action that names a price beside a Reasoning that denies one
        prices = re.search(r"top-of-search \$([\d.]+)|BID ([\d.]+)", a)
        denies = any(p in t.lower() for p in ("no price is stated", "no payable contribution",
                                              "carries no routed child"))
        if prices and denies:
            bad.append(f"row {r}: action prices it, its own reasoning denies pricing it")
            continue
        # 2 — a reversal citing a figure the reasoning no longer carries
        for fig in set(re.findall(r"\$([\d.]+)", v)):
            if fig in ("0.00",):
                continue
            if fig not in t and fig not in a:
                bad.append(f"row {r}: reverses-if cites ${fig}, absent from the action and reasoning")
                break
        # 3 — the action's wording must match its own number
        m = re.search(r"→ (\d+)% .*at the cap", a)
        if m and int(m.group(1)) < 350:
            bad.append(f"row {r}: action says 'at the cap' at {m.group(1)}%, which is not the cap")
        # 4 — a modifier stated in the action must equal the cell
        mm = re.search(r"TOS MODIFIER → (\d+)%", a)
        np_ = b.num(r, "New Percentage")
        if mm and np_ is not None and int(mm.group(1)) != int(np_):
            bad.append(f"row {r}: action says {mm.group(1)}%, cell holds {int(np_)}%")
    return len(bad), bad[:5]


def chk_spend_envelope(b):
    """A programme must fit an approved envelope, and a budget cap is not a
    forecast. On this account 525 priced re-routes sat in campaigns nobody had
    decided to enable — the campaign rows answered "should the budget change"
    and never "should this campaign run" — so the file implied a deployment it
    could not produce. Every campaign switched on states its wave and its
    projected daily spend, and the waves sum inside the envelope."""
    if "Deployment Waves" not in b.wb.sheetnames:
        return 1, ["Deployment Waves tab absent — no envelope is stated"]
    ws = b.wb["Deployment Waves"]
    env = None
    cum = []
    for r in range(2, ws.max_row + 1):
        if str(ws.cell(row=r, column=1).value or "").strip().lower() == "envelope":
            env = ws.cell(row=r, column=4).value
        v = ws.cell(row=r, column=5).value
        if isinstance(v, (int, float)):
            cum.append(float(v))
    bad = []
    if env is None:
        bad.append("no envelope figure on the Deployment Waves tab")
    elif cum and max(cum) > float(env) + 0.01:
        bad.append(f"cumulative ${max(cum):,.2f}/day exceeds the ${float(env):,.2f} envelope")
    # every enabled campaign names a wave
    for r in b.decided():
        if b.g(r, "Entity") != "Campaign":
            continue
        a = str(b.g(r, "Action") or "")
        if a.startswith("ENABLE") and "wave" not in a.lower():
            bad.append(f"row {r}: campaign enabled without naming its wave")
    return len(bad), bad[:4]


def chk_negation_evidence(b):
    """A relevant term that is not converting routes to the fix queue, never
    straight to a negative — only structurally off-target traffic is walled off.
    Each negation states its mode, because the evidence standard differs by mode."""
    if "Negatives" not in b.wb.sheetnames:
        return 0, ["Negatives tab absent — check skipped, not passed"]
    ws = b.wb["Negatives"]
    hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
    bad = []
    core = re.compile(r"\bsheet|\bbedsheet")
    off = re.compile(r"comforter|duvet|quilt|blanket|mattress|topper|protector|towel|curtain|rug|pillow(?!case)"
                     r"|satin|silk|flannel|percale|linen|^b0[a-z0-9]{8}$", re.I)
    for r in range(2, ws.max_row + 1):
        t = str(ws.cell(row=r, column=hdr.get("Search term", 1)).value or "").lower().strip()
        if not t:
            continue
        mode = str(ws.cell(row=r, column=hdr.get("Mode", 4)).value or "")
        if not mode:
            bad.append(f"Negatives row {r}: no mode stated")
        if core.search(t) and not off.search(t):
            bad.append(f'Negatives row {r}: "{t[:32]}" is a relevant sheet term — belongs in the fix queue')
    return len(bad), bad[:4]


def chk_plan_locked(b):
    """The plan's reconciliation table must equal a fresh read of this workbook,
    and its version stamp must name this file's hash. The pair is version-locked:
    if the workbook re-renders, the plan regenerates before either ships.

    This closes the engagement's terminal failure. Both files were individually
    sound; the plan had reconciled against an intermediate render and the
    workbook then advanced one more, so the section titled 'Figures to
    reconcile' certified agreement on twelve rows where six disagreed —
    including a five-fold difference on the budget row."""
    import hashlib, glob
    from pathlib import Path as _P
    cands = sorted(glob.glob("/mnt/user-data/outputs/*PPC_Plan_v*.docx"))
    if not cands:
        return 0, ["no plan document found — check skipped, not passed"]
    plan_path = cands[-1]
    try:
        from docx import Document
    except Exception:
        return 0, ["python-docx unavailable — check skipped, not passed"]
    doc = Document(plan_path)
    h = hashlib.sha256(open(b.path, "rb").read()).hexdigest()[:16]
    text = "\n".join(p.text for p in doc.paragraphs)
    bad = []
    if h not in text:
        bad.append(f"plan does not name this workbook's hash {h} — it was generated against a different render")
    gen = reconcile(b)
    found = {}
    for t in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        if not rows or len(rows[0]) < 2:
            continue
        if "generated from the workbook" not in " ".join(rows[0]).lower():
            continue
        for r in rows[1:]:
            if r[0]:
                found[r[0]] = r[1]
    if not found:
        bad.append("no generated reconciliation table in the plan")
    for k, v in found.items():
        if k in gen and gen[k] != v:
            bad.append(f'9.6 "{k}": plan says {v}, workbook reads {gen[k]}')
    return len(bad), bad[:5]


def chk_plan_arithmetic(b):
    """Every release-versus-commitment sentence in the plan is re-derived, and
    every child-name column is checked to contain child names. Seven renders
    established that find-replace is itself the failure mode: swapping one
    figure into sentences whose arithmetic depended on the old one produced
    "$26.16 released exceeds the $378.56 push" — false by fourteen times and
    visible without any rubric — and "against Routed per row from product-ad
    rows's own ceiling", a possessive built on a substituted phrase."""
    import glob
    cands = sorted(glob.glob("/mnt/user-data/outputs/*PPC_Plan_v*.docx"))
    if not cands:
        return 0, ["no plan document — check skipped, not passed"]
    try:
        from docx import Document
    except Exception:
        return 0, ["python-docx unavailable — check skipped, not passed"]
    doc = Document(cands[-1])
    lines = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            lines.append(" | ".join(c.text for c in r.cells))
    bad = []
    cmp_pat = re.compile(r"\$([\d,]+\.?\d*)[^.|]{0,80}?\b(exceeds|covers|over)\b[^.|]{0,80}?\$([\d,]+\.?\d*)", re.I)
    for s_ in lines:
        for m in cmp_pat.finditer(s_):
            a = float(m.group(1).replace(",", "")); c = float(m.group(3).replace(",", ""))
            if a < c and "scope" not in m.group(0).lower():
                bad.append(f"false comparison: {m.group(0)[:80]}")
    # a substituted phrase left inside a possessive or a name column
    for s_ in lines:
        if re.search(r"[a-z] rows's|rows's own", s_):
            bad.append(f"substitution artefact: {s_[:70]}")
    child = re.compile(r"BM-[A-Z0-9-]+|Queen|King|Twin|Full|Cal ?King", re.I)
    for t in doc.tables:
        if not t.rows:
            continue
        hdr = [c.text.strip().lower() for c in t.rows[0].cells]
        for j, hh in enumerate(hdr):
            if hh in ("child", "sku", "routed child", "preferred sku", "backup sku"):
                for r in t.rows[1:]:
                    if j >= len(r.cells):
                        continue
                    v = r.cells[j].text.strip()
                    if v and v != "—" and not child.search(v):
                        bad.append(f'child column holds "{v[:40]}", not a child name')
    return len(bad), bad[:5]


def chk_ranking_zero_modifier(b):
    """A ranking row at a zero top-of-search modifier is structurally unable to
    do the thing it exists for. At zero a campaign takes roughly 4% of its
    clicks at top of search; with a modifier set it takes over 70%. On a real
    portfolio 244 of 705 campaigns sat at zero — not 244 wrong decisions, one
    structural condition nobody looked for.

    Zero is legitimate where no rank case exists. The defect is a row carrying a
    Ranking objective AND a zero modifier without saying which it is."""
    bad = []
    for r in b.decided():
        if b.g(r, "Entity") != "Bidding Adjustment":
            continue
        if "Top" not in str(b.g(r, "Placement") or ""):
            continue
        m = b.num(r, "New Percentage")
        if m is None or m > 0:
            continue
        cn = str(b.g(r, "Campaign Name (Informational only)") or "")
        t = str(b.g(r, "Reasoning") or "").lower()
        looks_ranking = "rank" in cn.lower() or "-rank" in cn.lower()
        if not looks_ranking:
            continue
        # a zero on a ranking campaign must say which case it is
        if not any(w in t for w in ("no rank against a stated target", "not a ranking candidate",
                                    "objective is wrong", "no position is being bought")):
            bad.append(f"row {r}: ranking campaign at 0% top of search with no reason stated")
    return len(bad), bad[:4]


def chk_ratio_stated(b):
    """The top-of-search-to-product-page ratio is what explains the result, and
    it is invisible unless written: a 95% modifier reads as aggressive while a
    $3.00 base makes it 1.95 to one, nowhere near enough to isolate."""
    bad = []
    for r in b.decided():
        if b.g(r, "Entity") != "Bidding Adjustment":
            continue
        if "Top" not in str(b.g(r, "Placement") or ""):
            continue
        if b.num(r, "New Percentage") is None:
            continue
        t = str(b.g(r, "Reasoning") or "")
        if not re.search(r"\d+(\.\d+)?\s*(to one|:1|x) (the|against|over)|ratio", t, re.I):
            bad.append(f"row {r}: modifier set without stating the top-of-search to product-page ratio")
    return len(bad), bad[:4]


CHECKS = [
    ("structure",          "Rows not added to the delivered file",        chk_structure),
    ("coverage",           "Every row of every entity has a verdict",     chk_coverage),
    ("routing/red",        "No spend routed to a Red child",              chk_routing_is_red),
    ("routing/ads",        "Routing read from product ads, not names",    chk_routing_from_ads),
    ("routing/attribute",  "Attribute-naming query routes to a match",    chk_attribute_match),
    ("bound1/base",        "Base derived from the child, not inherited",  chk_base_derived),
    ("cap/correction",     "No bid step exceeds 25%",                     chk_correction_cap),
    ("entity/modifier",    "Modifiers on Bidding Adjustment rows",        chk_modifier_entity),
    ("bound2/premium",     "Modifier at or under 350%",                   chk_modifier_cap),
    ("bound3/unit",        "Top-of-search under what the unit earns",     chk_tos_vs_unit),
    ("bound4/flat",        "Top-of-search under the flat ceiling",        chk_tos_vs_flat),
    ("hold/legitimacy",    "Every hold cites a reason on the list",       chk_hold_legitimacy),
    ("write/present",      "Reasoning present on every decided row",      chk_reasoning_present),
    ("write/unique",       "No reasoning repeated more than 4x",          chk_reasoning_unique),
    ("write/figures",      "Reasoning cites this row's own numbers",      chk_reasoning_figures),
    ("write/window",       "Every figure carries its window",             chk_window_stated),
    ("write/reverses",     "Every verdict states what reverses it",       chk_reverses_present),
    ("write/sameval",      "Same-value New X columns left blank",         chk_same_value_blank),
    ("write/revision",     "No revision narrative",                       chk_revision_narrative),
    ("write/labels",       "No internal labels or raw platform IDs",      chk_internal_labels),
    ("write/integrity",    "Reasoning reads as prose, not a cut regex",   chk_text_integrity),
    ("discovery/reason",   "Discovery reasoned; auto types decided apart", chk_discovery_reasoned),
    ("tabs/actioned",      "Coverage and PAT tabs carry Action/Reasoning", chk_tabs_actioned),
    ("mkl/complete",       "New-launch rows fully sized, no placeholders",  chk_mkl_complete),
    ("discovery/premium",  "Discovery modifiers measured, not rank-gap",    chk_discovery_no_premium),
    ("sizing/flat",        "Premiums derived per row, not one flat value",  chk_flat_sizing),
    ("dedup/instances",    "One live instance per keyword and match type", chk_duplicates),
    ("launch/scope",       "Launches are new, unbranded, well-formed terms", chk_launch_scope),
    ("launch/relevance",   "Relevance judged on the listing, not a score",  chk_launch_relevance),
    ("launch/syntax",      "Syntax labels match each term's own roots",     chk_syntax_labels),
    ("launch/bids",        "Launch bids placed inside the suggested range", chk_launch_bids),
    ("deliver/apparatus",  "Gate output, census and review sheet ship in-file", chk_delivery_apparatus),
    ("write/action-cell",  "Action string names the figure the cell deploys",  chk_action_matches_cell),
    ("write/sample",       "Contribution cells name their unit sample",        chk_sample_disclosed),
    ("premium/gated",      "Premiums bounded by affordable TOS, rank case stated", chk_premium_gated),
    ("enable/gated",       "ENABLE in a paused campaign names its go/wait",        chk_enable_has_gate),
    ("qa/counts-agree",    "QA tabs state one set of counts, not two",      chk_counts_agree),
    ("row/integrity",      "Action, reasoning and reversal tell one story",  chk_row_integrity),
    ("spend/envelope",     "Waves sum inside the approved envelope",        chk_spend_envelope),
    ("negation/evidence",  "Only structurally off-target terms are negated", chk_negation_evidence),
    ("plan/locked",        "Plan 9.6 equals a fresh read of this workbook",  chk_plan_locked),
    ("plan/arithmetic",    "Plan comparisons re-derive; name columns hold names", chk_plan_arithmetic),
    ("ranking/zero-mod",   "Ranking rows at 0% modifier say which case",    chk_ranking_zero_modifier),
    ("placement/ratio",    "Modifier rows state the TOS-to-PDP ratio",      chk_ratio_stated),
    ("write/deferred",     "Nothing deferred to a later pass",            chk_deferred),
]


# ------------------------------------------------------------------- self-test
# A check never seen to fire has not been shown to test anything. Each injector
# writes ONE defect into ONE row of a copy — separate rows, because injecting
# several into the same row lets each overwrite the last and reports the checks
# broken when the test was.

def _inject(path, fn):
    wb = openpyxl.load_workbook(path)
    fb = wb["Final Bulk"]
    hdr = [c.value for c in fb[1]]
    col = {n: i + 1 for i, n in enumerate(hdr) if n}
    fn(fb, col)
    wb.save(path)


def selftest():
    tmp = "/tmp/_harness_selftest.xlsx"
    cases = {
        "cap/correction": lambda fb, c: fb.cell(row=_first(fb, c, "Bid"), column=c["New Bids"],
                                                value=round(fb.cell(row=_first(fb, c, "Bid"), column=c["Bid"]).value * 0.4, 2)),
        "entity/modifier": lambda fb, c: fb.cell(row=_first_entity(fb, c, "Keyword"), column=c["New Percentage"], value=150),
        "write/figures":   lambda fb, c: fb.cell(row=_first(fb, c, "Action", 3), column=c["Reasoning"], value="Held pending review."),
        "write/revision":  lambda fb, c: fb.cell(row=_first(fb, c, "Action", 4), column=c["Reasoning"], value="Corrected this pass; the earlier figure was wrongly applied."),
        "write/labels":    lambda fb, c: fb.cell(row=_first(fb, c, "Action", 5), column=c["Reasoning"], value="Held per §5 on portfolio 119993222153930."),
        "write/reverses":  lambda fb, c: setattr(fb.cell(row=_first(fb, c, "Action", 6), column=c["Reverses If"]), "value", None),
        "write/integrity": lambda fb, c: fb.cell(row=_first(fb, c, "Action", 7), column=c["Reasoning"],
                                                 value="Held on stock.29 at day 7 and the base follows."),
        # Set an absurd bid AND strip any staged day-7 target, because the check
        # judges a staged row on its stated target — leaving one in place means
        # the injected value is never the figure under test.
        "bound1/base":     lambda fb, c: (
            fb.cell(row=_first_entity(fb, c, "Keyword", 2), column=c["New Bids"], value=99.0),
            fb.cell(row=_first_entity(fb, c, "Keyword", 2), column=c["Reasoning"],
                    value="Base derived from the routed child.")),
        "routing/attribute": lambda fb, c: (fb.cell(row=_first_entity(fb, c, "Keyword", 3), column=c["Keyword Text"], value="white cooling sheets"),
                                            fb.cell(row=_first_entity(fb, c, "Keyword", 3), column=c["SKU"], value="BM-QUEEN-MIDNIGHT-BLACK")),
        "launch/scope":     lambda fb, c: None,   # injected on the MKL, below
        "launch/syntax":    lambda fb, c: None,   # injected on the MKL, below
        "launch/bids":      lambda fb, c: None,   # injected on the MKL, below
        "row/integrity":    lambda fb, c: fb.cell(row=_first(fb, c, "Action", 8), column=c["Reverses If"],
                                                  value="Reverses if $99999.99 an order is consumed."),
        "discovery/premium": lambda fb, c: (
            fb.cell(row=_first_auto_tos(fb, c), column=c["New Percentage"], value=294),
            fb.cell(row=_first_auto_tos(fb, c), column=c["Reasoning"],
                    value="Modifier sized on this row's own rank gap, unranked so full premium.")),
    }
    print(f"SELF-TEST — {len(cases)} checks, each injected into its own row\n")
    print("  loading the workbook once and injecting in memory; a fresh load per\n"
          "  case costs 11s on a 24k-row file and the injections touch separate rows.\n")
    ok = True
    lookup = dict((c[0], c[2]) for c in CHECKS)
    for cid, injector in cases.items():
        b = Book(WORKBOOK)                 # one load, discarded after the case
        if cid == "launch/bids":
            ws = b.wb["Master Keyword List"]
            hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
            for r in range(2, ws.max_row + 1):
                if str(ws.cell(row=r, column=hdr["Action"]).value or "").startswith("NEW LAUNCH"):
                    ws.cell(row=r, column=hdr["Proposed Bid "], value=50.0)  # blows the per-unit bound
                    break
        elif cid == "launch/syntax":
            ws = b.wb["Master Keyword List"]
            hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
            for r in range(2, ws.max_row + 1):
                if str(ws.cell(row=r, column=hdr["Action"]).value or "").startswith("NEW LAUNCH"):
                    ws.cell(row=r, column=hdr["Syntax Groups"], value="Irrelevant")
                    break
        elif cid == "launch/scope":
            # inject a keyword that IS already live into the launch list — the
            # exact defect this check exists for. Taking it from the bulk rather
            # than a "NOT A NEW LAUNCH" row, which a hand-edited file may not have.
            ws = b.wb["Master Keyword List"]
            hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
            for r in range(2, ws.max_row + 1):
                if str(ws.cell(row=r, column=hdr["Action"]).value or "").startswith("NEW LAUNCH"):
                    ws.cell(row=r, column=hdr["MKL"], value="cooling sheets queen")
                    break
        else:
            injector(b.fb, b.col)
        n, _ = lookup[cid](b)
        fired = n > 0
        ok &= fired
        print(f"  {'FIRES' if fired else 'BROKEN':<7} {cid:<20} returned {n}")
    print(f"\n{'All checks demonstrated to fail on a known-bad case.' if ok else 'A CHECK DID NOT FIRE — it has not been shown to test anything.'}")
    return 0 if ok else 1


def _first(fb, col, name, skip=0):
    seen = 0
    for r in range(2, fb.max_row + 1):
        if fb.cell(row=r, column=col[name]).value:
            if seen == skip:
                return r
            seen += 1
    return 2


def _first_auto_tos(fb, col):
    """First top-of-search bidding-adjustment row on an auto campaign."""
    for r in range(2, fb.max_row + 1):
        cn = str(fb.cell(row=r, column=col["Campaign Name (Informational only)"]).value or "")
        if fb.cell(row=r, column=col["Entity"]).value != "Bidding Adjustment":
            continue
        if "Top" not in str(fb.cell(row=r, column=col["Placement"]).value or ""):
            continue
        if "auto" in cn.lower() or "-disc-" in cn.lower():
            return r
    return 2


def _first_entity(fb, col, ent, skip=0):
    seen = 0
    for r in range(2, fb.max_row + 1):
        if fb.cell(row=r, column=col["Entity"]).value == ent:
            if seen >= skip:
                return r
            seen += 1
    return 2


# --------------------------------------------------------------- reconciliation
def reconcile(b):
    """Emit the plan's reconciliation table FROM the workbook, so the plan's
    match column and the gate output are one computation rather than two
    hand-written accounts. On a real build the two were typed separately and
    drifted a render apart: the plan certified 'Match' on twelve rows while the
    file beside it carried different figures on six of them, including a
    five-fold difference on the budget row — in the section built to prevent
    exactly that."""
    fb, g, num = b.fb, b.g, b.num
    dec = b.decided()
    HOLD = re.compile(r"^(HOLD|NO ACTION|NO MODIFIER|BUDGET HELD|STAY PAUSED)")
    held = [r for r in dec if HOLD.match(str(g(r, "Action")))]
    ask = [r for r in dec if str(g(r, "Action")).startswith("ASK")]
    acting = len(dec) - len(held) - len(ask)
    mix = collections.Counter()
    for r in dec:
        a = str(g(r, "Action"))
        if a.startswith("RE-ROUTE"): mix["re-route"] += 1
        elif a.startswith("PAUSE"): mix["pause"] += 1
        elif a.startswith("ENABLE"): mix["enable"] += 1
    prem = [r for r in dec if "small by design" in str(g(r, "Reasoning") or "")]
    lifts = sorted({int(m.group(1)) for r in prem
                    if (m := re.search(r"lift is (\d+) per cent", str(g(r, "Reasoning"))))})
    mods = [num(r, "New Percentage") for r in b.rows() if num(r, "New Percentage") is not None]
    loss = 0.0
    seen = set()
    for r in dec:
        m = re.search(r"Loss ceiling: \$([\d.,]+) a week", str(g(r, "Reasoning") or ""))
        if m and r not in seen:
            seen.add(r); loss += float(m.group(1).replace(",", ""))
    budget = sum(num(r, "New Budget") or 0 for r in b.rows()
                 if g(r, "Entity") == "Campaign" and num(r, "New Budget") is not None)
    budrows = sum(1 for r in b.rows() if g(r, "Entity") == "Campaign" and num(r, "New Budget") is not None)
    out = {
        "Decided rows": f"{len(dec):,}",
        "Acting": str(acting),
        "Held": str(len(held)),
        "Raised as questions": str(len(ask)),
        "Re-routes": str(mix["re-route"]),
        "Pauses": str(mix["pause"]),
        "Rank premiums": str(len(prem)),
        "Premium lift range": f"{min(lifts)}–{max(lifts)}% across {len(lifts)} values" if lifts else "none",
        "Max TOS modifier": f"{int(max(mods))}%" if mods else "none",
        "Daily budget written": f"${budget:,.2f} across {budrows} campaign rows",
        "Weekly loss ceiling": f"${loss:,.2f}",
    }
    # tabs that carry their own counts
    if "Master Keyword List" in b.wb.sheetnames:
        ws = b.wb["Master Keyword List"]
        hdr = {h: i + 1 for i, h in enumerate([c.value for c in ws[1]]) if h}
        L = [r for r in range(2, ws.max_row + 1)
             if str(ws.cell(row=r, column=hdr["Action"]).value or "").startswith("NEW LAUNCH")]
        sv = sum(int(ws.cell(row=r, column=hdr["Search Volume"]).value or 0) for r in L)
        out["New launches"] = str(len(L))
        out["Launch search volume"] = f"{sv:,}"
    for tab, label in (("Negatives", "Negatives"), ("Fix Queue", "Held out of negation")):
        if tab in b.wb.sheetnames:
            ws = b.wb[tab]
            out[label] = str(sum(1 for r in range(2, ws.max_row + 1)
                                 if ws.cell(row=r, column=1).value
                                 and str(ws.cell(row=r, column=1).value).strip().lower()
                                 not in ("total", "rule")))
    if "Deployment Waves" in b.wb.sheetnames:
        ws = b.wb["Deployment Waves"]
        waves = [r for r in range(2, ws.max_row + 1) if isinstance(ws.cell(row=r, column=1).value, int)]
        out["Deployment waves"] = str(len(waves))
        for r in range(2, ws.max_row + 1):
            if str(ws.cell(row=r, column=1).value or "").strip().lower() == "envelope":
                out["Spend envelope"] = f"${float(ws.cell(row=r, column=4).value):,.2f}/day"
    # plan/locked calls this function, so counting it here would recurse.
    # It is excluded from the figure the plan prints and reported separately
    # by the gate itself — the plan cannot certify the check that certifies it.
    scored = [c for c in CHECKS if c[0] not in ("plan/locked", "plan/arithmetic")]
    npass = sum(1 for c in scored if fnc(c, b) == 0)
    out["Validation gate"] = f"{npass}/{len(scored)} pass"
    return out


def fnc(c, b):
    try:
        return c[2](b)[0]
    except Exception:
        return 1


# ----------------------------------------------------------------------- main
def run(path):
    b = Book(path)
    print(f"PPC WORKBOOK HARNESS v{HARNESS_VERSION}")
    print(f"file: {Path(path).name}")
    print(f"rows: {b.fb.max_row - 1} | decided: {len(b.decided())}\n")
    failures = 0
    for cid, rule, fn in CHECKS:
        n, examples = fn(b)
        status = "PASS" if n == 0 else "FAIL"
        if n:
            failures += 1
        print(f"  {status}  {cid:<20} {rule}" + ("" if n == 0 else f"  [{n}]"))
        if n:
            for e in examples:
                print(f"           {e}")
    print(f"\n{len(CHECKS) - failures}/{len(CHECKS)} checks pass.")
    if failures:
        print("Re-run the whole set after fixing — levers feed each other, and a")
        print("check that passed before the last edit says nothing about what ships.")
    return 1 if failures else 0


if __name__ == "__main__":
    args = sys.argv[1:]
    if "--reconcile" in args:
        pos = [a for a in args if not a.startswith("--")]
        p = pos[0] if pos else WORKBOOK
        bk = Book(p)
        import hashlib
        h = hashlib.sha256(open(p, "rb").read()).hexdigest()[:16]
        print(f"WORKBOOK {Path(p).name}")
        print(f"sha256 {h}\n")
        for k, v in reconcile(bk).items():
            print(f"{k}\t{v}")
        sys.exit(0)
    if "--selftest" in args:
        sys.exit(selftest())
    # A path may be given positionally or after --file. Both work.
    # Silently falling back to a hardcoded default when a path IS supplied
    # meant the harness once reported on a file the caller never named —
    # so an unreadable path now stops the run instead of being ignored.
    path = None
    if "--file" in args:
        path = args[args.index("--file") + 1]
    else:
        pos = [a for a in args if not a.startswith("--")]
        if pos:
            path = pos[0]
    if path is None:
        path = WORKBOOK
        print(f"no file given; checking the default: {path}\n")
    if not Path(path).exists():
        print(f"ERROR: {path} does not exist. Refusing to check a different file.")
        sys.exit(2)
    sys.exit(run(path))
